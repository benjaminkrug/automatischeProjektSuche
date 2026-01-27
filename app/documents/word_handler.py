"""Word document handling with python-docx."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def fill_anschreiben(
    template_path: Path | str,
    output_path: Path | str,
    data: dict
) -> Path:
    """Fill Anschreiben template with project data.

    Placeholders in template:
    - {{DATUM}} - Current date
    - {{AUFTRAGGEBER}} - Client name
    - {{PROJEKTTITEL}} - Project title
    - {{KANDIDAT_NAME}} - Team member name
    - {{KANDIDAT_ROLLE}} - Team member role
    - {{PROJEKTBESCHREIBUNG}} - Brief project description
    - {{STAERKEN}} - Matching strengths
    - {{STUNDENSATZ}} - Proposed hourly rate

    Args:
        template_path: Path to Word template
        output_path: Path for output document
        data: Dictionary with placeholder values

    Returns:
        Path to created document
    """
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, data)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_placeholders_in_paragraph(paragraph, data)

    # Replace in headers/footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, data)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path


def _replace_placeholders_in_paragraph(paragraph, data: dict):
    """Replace all placeholders in a paragraph while preserving formatting."""
    full_text = paragraph.text

    # Check if any placeholder exists
    has_placeholder = False
    for key in data.keys():
        if f"{{{{{key}}}}}" in full_text:
            has_placeholder = True
            break

    if not has_placeholder:
        return

    # Replace placeholders
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value or ""))

    # Clear and rewrite paragraph
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)


def create_template_if_missing(template_path: Path | str) -> Path:
    """Create default Anschreiben template if it doesn't exist.

    Args:
        template_path: Path where template should be

    Returns:
        Path to template
    """
    template_path = Path(template_path)

    if template_path.exists():
        return template_path

    template_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set document properties
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Header with date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{DATUM}}")

    doc.add_paragraph()

    # Recipient
    p = doc.add_paragraph()
    p.add_run("{{AUFTRAGGEBER}}").bold = True

    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    p.add_run("Betreff: Bewerbung für \"{{PROJEKTTITEL}}\"").bold = True

    doc.add_paragraph()

    # Salutation
    doc.add_paragraph("Sehr geehrte Damen und Herren,")

    doc.add_paragraph()

    # Introduction
    p = doc.add_paragraph()
    p.add_run("mit großem Interesse habe ich Ihre Ausschreibung \"")
    p.add_run("{{PROJEKTTITEL}}").italic = True
    p.add_run("\" gelesen und bewerbe mich hiermit als ")
    p.add_run("{{KANDIDAT_ROLLE}}")
    p.add_run(".")

    doc.add_paragraph()

    # Project understanding
    doc.add_paragraph("{{PROJEKTBESCHREIBUNG}}")

    doc.add_paragraph()

    # Qualifications
    p = doc.add_paragraph()
    p.add_run("Meine Qualifikationen für dieses Projekt:").bold = True

    doc.add_paragraph("{{STAERKEN}}")

    doc.add_paragraph()

    # Rate
    p = doc.add_paragraph()
    p.add_run("Mein Stundensatz für dieses Projekt beträgt ")
    p.add_run("{{STUNDENSATZ}} €/Stunde").bold = True
    p.add_run(" zzgl. MwSt.")

    doc.add_paragraph()

    # Closing
    doc.add_paragraph(
        "Gerne stehe ich Ihnen für ein persönliches Gespräch zur Verfügung, "
        "um meine Eignung für dieses Projekt näher zu erläutern."
    )

    doc.add_paragraph()

    doc.add_paragraph("Mit freundlichen Grüßen,")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("{{KANDIDAT_NAME}}").bold = True

    # Save template
    doc.save(template_path)

    return template_path
